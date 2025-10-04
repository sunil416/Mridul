from utils.Document import Document
from docx import Document as DocxDocument


class LoadDOCX:
    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self) -> list:
        docx = DocxDocument(self.file_path)
        documents = []

        for i, para in enumerate(docx.paragraphs):
            text = para.text.strip()
            if text:
                documents.append(Document(
                    page_content=text,
                    metadata={"source": self.file_path, "paragraph": i + 1}
                ))
        self._documents = documents
        return documents
    
